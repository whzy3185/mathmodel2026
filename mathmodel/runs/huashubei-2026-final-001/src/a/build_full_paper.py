from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


RUN = Path(__file__).resolve().parents[2]
OUT = RUN / "paper"
FIG = RUN / "outputs" / "figures_v2"
RESEARCH_FIG = RUN / "outputs" / "figures_research"
RESULTS = json.loads((RUN / "outputs" / "data" / "final_results.json").read_text(encoding="utf-8"))

TITLE = "周期截断下导电介质网络的几何判定、概率夹逼与成本优化"
OUTPUT_STEM = "2023213805_赵奕程_第1次模拟训练论文"
BODY_FONT = "SimSun"
HEAD_FONT = "SimHei"
LATIN_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"
INK = "000000"


def set_run_font(run, east=BODY_FONT, latin=LATIN_FONT, size=12, bold=False, italic=False, color=INK):
    run.font.name = latin
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), east)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, top=None, bottom=None, left=None, right=None):
    """Set explicit cell borders; omitted sides are removed."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if spec is None:
            node.set(qn("w:val"), "nil")
            node.set(qn("w:sz"), "0")
        else:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(spec))
            node.set(qn("w:color"), "000000")
            node.set(qn("w:space"), "0")


def apply_three_line_borders(table):
    """Chinese contest-paper three-line table: 1.5 pt / 0.75 pt / 1.5 pt."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=12, bottom=6)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=12)


def set_table_geometry(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths_cm) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        dxa = int(width / 2.54 * 1440)
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(dxa))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            dxa = int(widths_cm[idx] / 2.54 * 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(widths_cm[idx])
            set_cell_margins(cell)


class Paper:
    def __init__(self):
        self.doc = Document()
        self.md: list[str] = []
        self.fig_no = 0
        self.table_no = 0
        self.eq_no = 0
        self._setup()

    def _setup(self):
        doc = self.doc
        sec = doc.sections[0]
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.35)
        sec.bottom_margin = Cm(2.25)
        sec.left_margin = Cm(2.45)
        sec.right_margin = Cm(2.35)
        sec.header_distance = Cm(1.2)
        sec.footer_distance = Cm(1.25)

        normal = doc.styles["Normal"]
        normal.font.name = LATIN_FONT
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        normal.font.size = Pt(11.5)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Cm(0.74)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.22

        for name, size, before, after in (("Heading 1", 15, 10, 5), ("Heading 2", 13, 7, 4), ("Heading 3", 12, 5, 3)):
            style = doc.styles[name]
            style.font.name = LATIN_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT)
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.first_line_indent = Cm(0)

        if "Figure Caption" not in [s.name for s in doc.styles]:
            cap = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        else:
            cap = doc.styles["Figure Caption"]
        cap.font.name = LATIN_FONT
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        cap.font.size = Pt(10)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(5)
        cap.paragraph_format.keep_with_next = False
        cap.paragraph_format.first_line_indent = Cm(0)

        footer = sec.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        self._append_field(p, "PAGE", "1")

    @staticmethod
    def _append_field(p, instruction: str, display: str):
        run = p.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = display
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for node in (begin, instr, separate, text, end):
            run._r.append(node)
        set_run_font(run, size=10)

    def meta_line(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("学号：2023213805    姓名：赵奕程    第1次模拟训练论文")
        set_run_font(r, size=10.5)

    def title(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        r = p.add_run(TITLE)
        set_run_font(r, east=HEAD_FONT, size=18, bold=True)
        self.md += [f"# {TITLE}", ""]

    def abstract_heading(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("摘要")
        set_run_font(r, east=HEAD_FONT, size=15, bold=True)
        self.md += ["## 摘要", ""]

    def paragraph(self, text: str, *, first=True, size=11.5, after=3, keep=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74) if first else Cm(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.22
        p.paragraph_format.keep_with_next = keep
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        set_run_font(r, size=size)
        self.md += [text, ""]
        return p

    def keyword(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(3)
        r = p.add_run("关键词：")
        set_run_font(r, east=HEAD_FONT, size=11.5, bold=True)
        r = p.add_run(text)
        set_run_font(r, size=11.5)
        self.md += [f"**关键词：** {text}", ""]

    def page_break(self):
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    def heading(self, text: str, level=1):
        p = self.doc.add_paragraph(style=f"Heading {level}")
        r = p.add_run(text)
        set_run_font(r, east=HEAD_FONT, size={1: 15, 2: 13, 3: 12}[level], bold=True)
        self.md += [f"{'#' * (level + 1)} {text}", ""]

    def equation(self, expression: str, explanation: str | None = None, md_expression: str | None = None):
        self.eq_no += 1
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_together = True
        p_pr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9000")
        tabs.append(tab)
        p_pr.append(tabs)
        math_para = OxmlElement("m:oMathPara")
        omath = OxmlElement("m:oMath")
        mr = OxmlElement("m:r")
        mt = OxmlElement("m:t")
        mt.text = expression
        mr.append(mt)
        omath.append(mr)
        math_para.append(omath)
        p._p.append(math_para)
        r = p.add_run(f"\t({self.eq_no})")
        set_run_font(r, latin=MATH_FONT, size=11.5)
        md_math = md_expression or expression
        self.md += [f"$$ {md_math} \\tag{{{self.eq_no}}} $$", ""]
        if explanation:
            self.paragraph(explanation)

    def figure(self, filename: str, caption: str, width_cm=15.5, *, research=False):
        self.fig_no += 1
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        figure_root = RESEARCH_FIG if research else FIG
        picture = p.add_run().add_picture(str(figure_root / filename), width=Cm(width_cm))
        picture._inline.docPr.set("title", f"图{self.fig_no} {caption}")
        picture._inline.docPr.set("descr", caption)
        cap = self.doc.add_paragraph(style="Figure Caption")
        cap.paragraph_format.keep_together = True
        r = cap.add_run(f"图{self.fig_no}  {caption}")
        set_run_font(r, size=10)
        rel = "figures_research" if research else "figures_v2"
        self.md += [f"![图{self.fig_no} {caption}](../outputs/{rel}/{filename})", "", f"图{self.fig_no}  {caption}", ""]

    def table(self, caption: str, headers, rows, widths_cm, aligns=None, font_size=9.5):
        self.table_no += 1
        cap = self.doc.add_paragraph(style="Figure Caption")
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(2)
        cap.paragraph_format.keep_with_next = True
        r = cap.add_run(f"表{self.table_no}  {caption}")
        set_run_font(r, size=10)
        table = self.doc.add_table(rows=1, cols=len(headers))
        # Keep the template's neutral table base; the three-line rule below is
        # applied explicitly and does not depend on localized style names.
        table.style = None
        set_table_geometry(table, widths_cm)
        set_repeat_table_header(table.rows[0])
        aligns = aligns or [WD_ALIGN_PARAGRAPH.CENTER] * len(headers)
        for idx, h in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(h))
            set_run_font(r, east=HEAD_FONT, size=font_size, bold=True)
        for row in rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cell = cells[idx]
                cell.text = ""
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                p.alignment = aligns[idx]
                r = p.add_run(str(value))
                set_run_font(r, size=font_size)
        set_table_geometry(table, widths_cm)
        apply_three_line_borders(table)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)
        spacer.paragraph_format.first_line_indent = Cm(0)

        self.md += [f"**表{self.table_no}  {caption}**", ""]
        self.md.append("| " + " | ".join(map(str, headers)) + " |")
        self.md.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            self.md.append("| " + " | ".join(str(v) for v in row) + " |")
        self.md.append("")

    def code_block(self, code: str):
        lines = code.splitlines()
        for index, line in enumerate(lines):
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = index < len(lines) - 1
            r = p.add_run(line or " ")
            set_run_font(r, east="Microsoft YaHei", latin="Consolas", size=7.4)
        self.md += ["```python", code, "```", ""]

    def save(self, docx: Path, md: Path):
        props = self.doc.core_properties
        props.title = TITLE
        props.subject = "2026华数杯A题第一次模拟训练论文"
        props.author = "赵奕程"
        props.last_modified_by = "赵奕程"
        self.doc.save(docx)
        md.write_text("\n".join(self.md), encoding="utf-8")


def q1_result_rows():
    rows = []
    for item in RESULTS["Q1"]:
        path = "-".join(str(x).replace("LEFT", "左面").replace("RIGHT", "右面") for x in item["conductive_path_1_based"]) if item["connected"] else "无"
        rows.append([item["group"], item["row_count"], item["edge_count"], item["left_contact_count"], item["right_contact_count"], "导通" if item["connected"] else "不导通", path])
    return rows


def q1_cert_rows():
    rows = []
    for item in RESULTS["Q1"]:
        for c in item["path_certificates"]:
            if c["type"] == "electrode_contact":
                rows.append([
                    item["group"], f"{c['from']}-{c['to']}", f"{c['surface_gap_nm']:.3f}",
                    "—", "—", "—", "—", "电极接触通过",
                ])
                continue
            if c["type"] == "interior_side_to_side":
                s, t = c["segment_parameters"]
                rows.append([
                    item["group"], f"{c['from']}-{c['to']}", f"{c['axis_distance_nm']:.3f}",
                    f"{s:.3f}", f"{t:.3f}", f"{c['minimum_endpoint_margin_nm']:.1f}",
                    f"{c['axis_connector_orthogonality_residual']:.1e}", "通过",
                ])
    return rows


def build_paper_legacy():
    OUT.mkdir(parents=True, exist_ok=True)
    p = Paper()
    p.meta_line()
    p.title()
    p.abstract_heading()
    p.paragraph("本文研究周期立方微构体中导电介质的左右导通判定、随机填充概率界与成本优化问题。针对附件给出的确定性构型，将每个平端圆柱视为图节点，以介质间最短距离和介质—电极距离为连边准则；胶囊体仅用于生成保守候选边，正例路径再由轴段内部最短点完成几何核验。随机填充部分首先声明独立均匀中心、独立各向同性取向以及周期平移片段保持母体导体身份三项概率模型条件；随后由单粒子直接贯通事件构造总导通概率下界，并对“无直接贯通但仍存在通路”的剩余事件建立终端薄壳联合上界。最后在整数域内枚举全部低成本方案，分别讨论允许某类介质为零和两类介质均为正两种口径。")
    p.paragraph("在上述概率模型条件下，附件组1不导通，组2和组3导通，显式通路分别为左面-2-12-24-39-右面和左面-63-264-216-351-右面。A 的体积分数为0.50%、0.60%、0.70%和1.00%时，仅直接贯通事件已使总不导通概率上界分别降至10^-45.20、10^-54.13、10^-63.20和10^-90.27量级；这些下界严格小于1，不能因浮点舍入写成精确等于1。仅填充A时，7根A的总导通概率上界为0.872279，8根A的直接贯通概率下界为0.904810，故最低数量为8根，对应体积分数0.01131%，按百分号后两位报告为0.01%。")
    p.paragraph("混合填充中，按题目“同时填充”的严格语义，1A+50B为两类数量均为正时的最低成本方案，成本0.09862元；放宽为非负整数域后，边界解0A+57B成本0.09550元。全文区分解析事件概率、总导通下界和总导通上界，阈值与最优性均由不可行侧的上界证据闭合。")
    p.keyword("周期边界；随机几何图；导通概率界；联合界；整数优化")
    p.page_break()

    p.heading("1  问题重述与问题分析")
    p.heading("1.1  问题背景与研究对象", 2)
    p.paragraph("题目给定边长L=10000 nm的立方微构体，左右带电面位于x=-5000 nm与x=5000 nm。当导电介质之间或介质与带电面的表面最短距离不超过g=1.8 nm时，视为接触导通。介质A为高度H=5000 nm、半径r_A=30 nm的平端直圆柱；介质B为半径r_B=200 nm的球。题面规定越界部分平移一个边长后从对侧进入。随机问题还需要补充解释：本文把这些平移片段视为母体导体的周期表示，电学身份保持不变。该解释是Q2-Q4解析结论成立的条件，图1概括几何对象和接触口径。")
    p.figure("S01_problem_geometry.png", "问题几何、介质类型与周期边界示意", 15.3, research=True)
    p.heading("1.2  四个问题的输入、输出与难点", 2)
    p.paragraph("问题一的输入是附件中的三组圆柱轴段坐标，输出是每组是否导通及可复核的通路；难点在于平端圆柱不能直接等同于端部为半球的胶囊体。问题二将确定性坐标改为随机位置和随机方向，要求计算四个体积分数下的导通概率。问题三寻找仅填充A且导通概率不低于90%的最低填充率，需要同时证明候选值充分和前一整数不足。问题四加入B及材料成本，形成带概率约束的二元整数优化，并存在“是否允许某类介质数量为零”的题意歧义。")
    p.heading("1.3  总体技术路线", 2)
    p.paragraph("本文把四问组织为三个层次：第一层为确定性几何图和路径证书；第二层为所有随机问题共享的直接贯通下界与非直接通路上界；第三层为基于概率界的整数阈值和成本优化。该结构使Q2-Q4共享同一母模型，避免每问重复定义概率口径。技术路线如图2所示。")
    p.figure("S02_workflow.png", "四问依赖关系与总体技术路线", 15.3, research=True)
    p.heading("1.4  基线方法与模型选择", 2)
    p.paragraph("最直接的基线是逐次随机生成全部介质、构造几何接触图，再以蒙特卡洛频率估计导通概率。这一方法适合估计一般构型的真实总导通概率，但阈值附近若要区分7根不足与8根充分，有限样本频率只能给统计区间，不能自动给出整数最小性的证明；问题四还需要对大量整数配方重复模拟。本文因此保留确定性几何图作为问题一的执行层，在问题二至四使用解析上下界作为主证明。解析界不是对全几何仿真的替代品，而是为阈值和全局最优性提供可核验的不可行证书。")
    p.table("候选建模路线及本文取舍", ["路线", "可回答内容", "主要不足", "本文用途"], [
        ["逐样本几何图+蒙特卡洛", "真实总导通概率的区间估计", "需通用周期平端实体距离核；阈值证明成本高", "未来校准"],
        ["连续渗流阈值", "大体系的簇形成与长径比效应", "与有限盒、电极和单粒子跨界规则不等价", "机制对照"],
        ["直接贯通解析下界", "候选方案的严格可行性", "不计多粒子曲折通路", "Q2、Q3、Q4充分性"],
        ["终端薄壳联合上界", "低填充方案的严格不足性", "通常偏松，不估计真实概率", "Q3、Q4排除证书"],
    ], [3.1, 4.5, 5.2, 4.1], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.7)

    p.heading("2  模型假设、符号与量纲")
    p.heading("2.1  题面条件与补充假设", 2)
    p.table("题面条件、补充假设及失效影响", ["类别", "内容", "使用位置", "改变后的影响"], [
        ["题面条件", "每行附件数据表示一个A；边界越界部分周期平移", "Q1-Q4", "改变介质身份或周期解释会改变全部结果"],
        ["几何口径", "Q1按附件每行一个A建图；不跨行合并介质身份", "Q1", "跨行合并会制造同体边"],
        ["补充假设", "介质中心独立且在立方体内均匀分布", "Q2-Q4", "排斥或团聚会破坏独立乘积式"],
        ["补充假设", "A的轴向独立且各向同性", "Q2-Q4", "取向偏置会改变q_A及整数阈值"],
        ["题面条件", "随机介质允许相互贯穿、重叠", "Q2-Q4", "若增加排斥约束则概率模型改变"],
        ["补充解释", "周期平移片段保持母体导体的电学身份", "Q2-Q4", "若片段电学断开，则直接贯通机制失效"],
    ], [2.2, 6.2, 2.4, 5.1], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9)
    p.heading("2.2  主要符号", 2)
    p.table("主要符号及含义", ["符号", "含义", "取值或单位"], [
        ["L", "立方体边长", "10000 nm"], ["g", "表面导通阈值", "1.8 nm"],
        ["H, r_A", "介质A的高度与半径", "5000 nm, 30 nm"], ["r_B", "介质B半径", "200 nm"],
        ["n_A, n_B", "A、B填充数量", "非负整数"], ["a_x", "粒子在x方向的支撑半宽", "nm"],
        ["D", "至少一个粒子直接贯通事件", "事件"], ["T", "整体左右导通事件", "事件"],
        ["S_i^L, S_i^R", "粒子i未越界但接触左/右电极的剩余薄壳事件", "事件"],
        ["P_dir", "直接贯通事件D的概率", "[0,1]"], ["C", "材料成本", "元"],
    ], [2.7, 8.7, 4.5], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER], font_size=9.4)
    p.heading("2.3  体积与成本换算", 2)
    p.equation("V_A=πr_A²H=1.4137167×10⁷ nm³,   V_B=4πr_B³/3=3.3510322×10⁷ nm³")
    p.equation("C=1.05V_An_A/10⁹+0.05V_Bn_B/10⁹=0.0148440253n_A+0.0016755161n_B")
    p.paragraph("式(1)按真实圆柱与球体积计算，式(2)把nm^3换算为um^3。所有数量先在整数域内求解，体积分数仅作为结果的物理表达，不用四舍五入后的体积分数反推粒子数。")

    p.heading("3  问题一：确定性构型的几何连通模型")
    p.heading("3.1  附件数据审计", 2)
    p.paragraph("三组附件分别包含12、49和535行介质A数据。轴段长度范围为1019.117-5000.000 nm、577.474-5000.000 nm和525.697-5000.000 nm，说明坐标中保留了边界截断后的轴段表达。组1、组2和组3触及任一边界面的行比例分别为58.3%、44.9%和71.0%。附件又明确每行表示一个A，因此本文不因相对边界端点重合而跨行合并身份，也不把短轴段擅自延长。数据特征见图3。")
    p.figure("S03_data_audit.png", "附件三组数据规模与边界截断特征", 15.3, research=True)
    p.paragraph("附件中还存在坐标数值相同或相差固定量的端点，但题面只提供每行两端坐标，没有原始中心、原始方向和周期片段ID。仅凭端点巧合无法唯一判定两行是否来自同一母体；特别是±500并不是边长10000 nm正方体的边界，不能据此合并。为保证结果可重现，主口径采用题面附件说明的行级身份。若未来取得片段ID，应把同一母体的所有片段并入一个图节点，并重新计算与电极及其他介质的最短距离；当前结论明确限定在可识别的行级轴段口径。")
    p.heading("3.2  行级连通图", 2)
    p.paragraph("建立无向图G=(V,E)。V由两个虚拟电极节点LEFT、RIGHT和每一行对应的圆柱节点组成。若圆柱与电极的最短距离不超过g，则连接圆柱节点与相应电极；若两个圆柱的表面最短距离不超过g，则连接对应节点。LEFT与RIGHT位于同一连通分量当且仅当该构型导通。")
    p.equation("(i,j)∈E ⇔ dist(A_i,A_j)≤g;   (LEFT,i)∈E ⇔ dist(A_i,F_L)≤g")
    p.paragraph("实现时先由轴段采样点的KD树生成保守候选对，再对候选轴段使用闭式最短距离核复核；附件规模最大仅535行，另以全对全O(n²)枚举作为独立对照。电极接触不只检查端点x坐标，还将圆柱端面圆盘在x方向的径向投影计入支撑范围。最后用广度优先搜索判定连通，并由父指针恢复一条显式LEFT—RIGHT路径。")
    p.code_block("输入：每行轴段端点 P_i,Q_i\n1. 计算每根圆柱对左右电极的表面间隙\n2. KD 广相位生成候选对；精确轴段距离筛选胶囊超图边\n3. 将 LEFT、RIGHT 与接触圆柱连边\n4. BFS 搜索 LEFT 到 RIGHT；若成功则回溯路径\n5. 对路径每条圆柱边计算 s、t、端面余量与正交残差\n输出：连通布尔值、显式路径、逐边证书")
    p.paragraph("KD广相位的作用只是减少需要精算的候选数，不能改变最终判定；全对全对照用于防止采样间距造成漏边。BFS复杂度为O(|V|+|E|)，最耗时部分是候选边距离核。所有比较均保留1e-10 nm级数值容差，容差只吸收浮点舍入，不扩大题设1.8 nm阈值。")
    p.heading("3.3  平端圆柱的候选边与充分证书", 2)
    p.paragraph("两条轴段的最短距离不超过2r_A+g=61.8 nm时，相应胶囊体必接触，因此该条件可用于生成候选边。但胶囊体把圆柱端部替换为半球，可能把端点附近接近误判为平端圆柱接触。为避免这一问题，对正例路径中的每条圆柱—圆柱边保存轴段最短点参数s,t。当0<s<1且0<t<1时，两个最短点均位于轴段内部，表面间隙等于max(0,d_axis-2r_A)，可作为真实侧面接触的充分证书。对负例，胶囊体是平端圆柱的超集；若胶囊超图仍不连通，则真实图一定不连通。图4解释了证书条件。")
    p.figure("S04_flat_cylinder_certificate.png", "平端圆柱侧面接触的轴段最短点证书", 14.7, research=True)
    p.paragraph("更具体地，设两轴段参数式为P(s)=P_0+s(P_1-P_0)、Q(t)=Q_0+t(Q_1-Q_0)。内部极小点满足连接向量P(s)-Q(t)同时垂直于两轴方向，因此沿该连接向量分别向两条轴外移r_A后，仍落在对应截面圆盘内，实体表面间隙恰为d_axis-2r_A。正文同时保存端面余量min{sL_i,(1-s)L_i,tL_j,(1-t)L_j}；其为正说明接触截面不落在端面退化区。正交残差则用于排查数值求解是否真的满足内部驻点条件。")
    p.paragraph("若s或t落在0、1附近，轴段距离减半径只能给胶囊体候选，真实平端圆柱可能属于侧面—端面、端面—端面或圆周—侧面情形，本文不把这类边直接当作正例证书。该不对称策略保证：组1用超集不连通证明负例不会误判；组2、组3只选取通过内部侧面条件的路径证明正例，也不会依赖未经认证的端部候选边。")
    p.heading("3.4  三组构型的求解结果", 2)
    p.paragraph("对每组数据先计算电极接触，再全对全生成候选边，以广相位空间索引复核边数，最后用广度优先搜索恢复LEFT-RIGHT路径。图5把三组真实xyz构型并列展示，红线只标注算法恢复的路径；三维图用于解释空间关系，结论仍以数值图搜索和路径证书为准。")
    p.figure("C03_q1_path_certificate.png", "组2与组3的三维导通路径及相邻杆件最近点证书", 15.8, research=True)
    p.table("Q1三组附件构型的连通结果", ["组别", "节点数", "边数", "左接触", "右接触", "结论", "显式路径"], q1_result_rows(), [1.1, 1.4, 1.3, 1.3, 1.3, 1.5, 8.1], aligns=[WD_ALIGN_PARAGRAPH.CENTER]*6+[WD_ALIGN_PARAGRAPH.LEFT], font_size=8.6)
    p.paragraph("组1在更易连通的胶囊超图中仅有2条介质间候选边，LEFT与RIGHT仍分属不同连通分量，因此按附件行级口径不导通；由于真实平端圆柱图是该超图的子图，这一负结论不依赖端部距离分类。组2和组3均找到显式路径。逐边证书表不仅列出6条介质间边的s、t、端面余量和正交残差，也列出两条路径各自的LEFT首边与末边RIGHT电极表面间隙，从而覆盖完整路径而不是只检查中间杆件。")
    p.table("Q1正例完整路径的逐边证书", ["组别", "边", "距离/nm", "s", "t", "端面余量/nm", "正交残差", "结论"], q1_cert_rows(), [1.1, 1.8, 2.0, 1.2, 1.2, 2.5, 2.3, 1.4], font_size=8.2)

    p.heading("4  Q2-Q4共享的随机填充概率模型")
    p.heading("4.1  单粒子直接贯通事件", 2)
    p.paragraph("若某粒子在x方向跨越任一周期边界，其越界片段从对侧进入且保持同一导体身份，于是该粒子自身形成左右贯通。该事件是总导通事件T的充分条件而非必要条件，因此由它得到的是严格概率下界。图7分别示意圆柱A和球B的直接贯通机制。")
    p.figure("S05_direct_bridge_mechanism.png", "A、B单粒子跨越周期边界的直接贯通机制", 15.3, research=True)
    p.heading("4.2  圆柱A的直接贯通概率", 2)
    p.paragraph("设圆柱轴向单位向量的x分量为u_x。平端圆柱在x方向的支撑半宽由轴向投影和端面圆盘投影共同组成。固定方向后，圆柱中心落入任一边界内侧a_x范围即越界，两侧合计条件概率为2a_x/L。")
    p.equation("a_x=(H/2)|u_x|+r_A√(1-u_x²)")
    p.paragraph("圆柱轴是无向的，u与-u代表同一取向。各向同性轴向可在半球上取均匀面测度；等价地，U=|u_x|在[0,1]上服从均匀分布。因此两个期望可以直接化为一维积分：E|u_x|=∫_0^1u du=1/2，E√(1-u_x²)=∫_0^1√(1-u²)du=π/4。由于本题最大支撑半宽加阈值仍小于L/2，左右两条越界中心区间互不重叠，条件概率2a_x/L无需再截断。")
    p.paragraph("将支撑半宽对取向积分即可得到单根A的直接贯通概率q_A。图中的曲线是固定取向条件概率，水平线是各向同性方向平均；二者不能混为一谈。")
    p.equation("q_A=2E(a_x)/L=(H/2+πr_A/2)/L=0.25471238898")
    p.figure("S06_orientation_support.png", "圆柱取向对支撑半宽和条件越界概率的影响", 15.3, research=True)
    p.heading("4.3  球B与多粒子的直接贯通下界", 2)
    p.paragraph("球的支撑半宽恒为r_B，故球心落入任一边界内侧r_B范围即可越界。不同粒子的位置和方向独立时，至少一个粒子直接贯通的概率可由补事件乘积得到。")
    p.equation("q_B=2r_B/L=0.04")
    p.equation("P_dir=P(D)=1-(1-q_A)ⁿᴬ(1-q_B)ⁿᴮ")
    p.heading("4.4  非直接通路的剩余薄壳联合上界", 2)
    p.paragraph("为证明某个较低填充方案不可能达到90%，仅有下界不够。定义D_i为粒子i直接越界，C_i^L为粒子i接触左电极。固定其x向支撑半宽a_i时，无条件接触层宽为a_i+g，因此不能把P(C_i^L)误写成g/L。关键在于研究无直接贯通事件D^c中的剩余接触：粒子既不越界又接触左电极时，中心只能位于宽度恰为g的薄壳。")
    p.equation("S_iᴸ=C_iᴸ∩D_iᶜ={-L/2+a_i<X_i≤-L/2+a_i+g},   P(S_iᴸ)=g/L")
    p.paragraph("令N=T交D^c。任何N中的左右路径都必须包含不同的终端粒子i和j，分别落入左、右剩余薄壳；若由同一粒子承担两端接触，则该粒子已经属于直接贯通事件D。对所有有序粒子对使用独立性和Boole联合界[3]，得到式(9)。事件关系见图9。")
    p.equation("P_dir≤P(T)≤min{1, P_dir+n(n-1)(g/L)²},   n=n_A+n_B")
    p.figure("S07_event_bounds.png", "总导通、直接贯通与剩余薄壳事件的上下界关系", 15.3, research=True)
    p.paragraph("式(9)可整理为一个终端薄壳上界命题。条件为：各粒子中心相互独立且x坐标均匀；单粒子支撑半宽加g小于L/2；周期越界片段保持同一介质身份。证明分四步：首先将总导通事件分解为D与N=T∩D的补集；其次，在N中任一左右通路都必须各有一个未越界的左、右终端；再次，同一未越界粒子因尺寸条件不可能同时承担两端，故终端是有序的不同粒子对；最后，对每一有序对使用跨粒子独立性得到(g/L)²，再对n(n−1)个有序对应用Boole联合界。中间链路被完全放宽为“一定存在”，所以该式宁可偏大也不会漏掉曲折通路。")
    p.paragraph("上界增量看似与粒子形状无关，是因为支撑半宽a_i对应的越界区间已经被事件D剔除；剩余的电极接触层只比越界区间多出固定宽度g。若误把无条件电极接触概率写成g/L会低估终端事件，而本文只在D的补集内使用该薄壳宽度。")
    p.heading("4.5  概率模型的解释边界", 2)
    p.paragraph("式(9)的上界用于证明“不足”，不用于精确估计真实导通概率；它忽略了中间粒子如何连接，因此通常偏松，但不会漏掉曲折通路。式(8)的下界用于证明“充分”，也不等于总导通概率。只有当上下界位于阈值两侧时，才能给出严格整数阈值或最优性结论。若粒子中心存在排斥、团聚或取向相关，式(8)中的独立乘积和式(9)中的跨粒子独立性需要重建。")
    p.heading("4.6  与连续渗流模型的关系", 2)
    p.paragraph("随机杆体系常用排除体积、随机几何图或连续渗流阈值刻画大体系中多粒子簇的形成[5-10]。这些研究说明边概率会同时受到体积分数、形状、取向分布、接触壳和分散状态影响；对允许重叠的随机取向胶囊体，有限尺寸模拟也是成熟路线[10]。但胶囊体端部为半球，本题为平端圆柱；同时本题是有限立方体、指定左右电极，并采用特殊的周期截断规则，单粒子跨界事件已提供可计算的充分条件。因此本文不把热力学极限阈值移植为答案，只把它用作机制和量级的外部参照。若后续需要估计上下界之间的真实总概率，应先实现完整平端圆柱周期距离核，再做有限盒蒙特卡洛，而不能用平均邻居数替代。")

    p.heading("5  问题二：给定体积分数下的导通概率")
    p.heading("5.1  体积分数到整数数量的换算", 2)
    p.paragraph("设目标体积分数为phi。粒子数必须取整数，本文选择使n_A V_A/L^3最接近目标phi的整数n_A，并同时报告目标值与实际值。四个目标体积分数对应354、424、495和707根A。")
    p.equation("n_A=round(φL³/V_A),   φ_actual=n_AV_A/L³")
    p.heading("5.2  概率结果与解释", 2)
    q2_rows = [[f"{100*r['requested_fraction']:.2f}%", r["a_count"], f"{100*r['achieved_fraction']:.5f}%", f"{r['log10_failure_probability_upper_bound']:.2f}", f"至少1-10^{r['log10_failure_probability_upper_bound']:.2f}"] for r in RESULTS["Q2"]]
    p.table("Q2体积分数、整数数量与直接贯通下界", ["目标体积分数", "A数量", "实际体积分数", "log10不导通上界", "导通概率下界"], q2_rows, [2.5, 1.7, 2.7, 4.0, 5.0], font_size=9)
    p.paragraph("结果表中的“不导通上界”仅考虑没有任何A直接贯通的概率(1-q_A)^n_A。粒子间进一步接触只会增加总导通概率，因此该量确为总不导通概率的上界。四个数量级远低于常规数值显示精度，故可以表述为“在报告精度内导通概率为1”，但不能写成数学上精确等于1。随后图形展示其对数数量级。")
    p.figure("C05_q2_failure_lollipop.png", "Q2四种体积分数下不导通概率上界的数量级", 15.3, research=True)
    p.heading("5.3  数值稳定表达与概率区间", 2)
    p.paragraph("当n_A达到数百时，直接计算1-(1-q_A)^n_A会在双精度浮点中舍入为1，使“极接近1”和“严格等于1”无法区分。本文先计算log10 P(T的补集)≤n_A log10(1-q_A)，把失败概率的数量级作为机器结果保存；论文再写成P(T)∈[1-10^k,1]，其中k为表中负数。这样既保留数学上的严格性，也避免把浮点下溢当成物理必然。")
    p.paragraph("四个工况的区间宽度从约10^-45降至10^-90。它们说明在本文概率生成机制下，单粒子跨界事件已经足以使失败风险极小；区间并不等于真实总导通概率的精确值，因为多粒子接触簇只会进一步抬高导通概率。若改变取向分布或取消周期身份，这些数量级必须重新计算，不能从体积分数单独外推。")

    p.heading("6  问题三：仅填充A的最低填充量")
    p.heading("6.1  候选阈值定位", 2)
    p.paragraph("由式(8)可知直接贯通下界随n_A单调增加。解1-(1-q_A)^n_A>=0.90可将候选值定位在8根。证明“8根是最小值”还需验证8根充分和7根不足。")
    p.equation("n_A≥⌈log(0.10)/log(1-q_A)⌉=8")
    p.heading("6.2  7根不足与8根充分", 2)
    p.paragraph("当n_A=8时，直接贯通概率下界为0.9048100243>0.90，故8根充分。当n_A=7时，直接贯通概率为0.8722775285；剩余非直接通路上界增量为7*6*(1.8/10000)^2=1.3608e-6，因而总导通概率上界为0.8722788893<0.90，故7根不足。图11显示上下界几乎重合但分别位于阈值两侧。")
    q3_rows = [[r["a_count"], f"{r['direct_bridge_lower_bound']:.6f}", f"{r['non_direct_path_upper_addition']:.2e}", f"{r['conduction_upper_bound']:.6f}"] for r in RESULTS["Q3"]["proof_rows"]]
    p.table("Q3从1至8根A的导通概率上下界", ["A数量", "直接贯通下界", "非直接上界增量", "总导通上界"], q3_rows, [2.3, 4.4, 4.7, 4.4], font_size=9.2)
    p.figure("C04_q3_bounds_band.png", "Q3中7根不足、8根充分的解析夹逼", 15.3, research=True)
    p.heading("6.3  单调耦合与整数最小性", 2)
    p.paragraph("为说明只检查7和8两点即可闭合最小性，把n根和n+1根的随机构型耦合在同一概率空间：先生成前n根，再独立加入第n+1根。新增节点和接触边不会删除原有LEFT—RIGHT路径，因此真实总导通事件T_n包含于T_{n+1}，总导通概率关于粒子数单调不减。于是7根上界低于0.90可同时排除0至7根，8根下界高于0.90则证明8根可行。这个结论针对真实总导通事件，而不只针对解析下界。")
    p.heading("6.4  填充率报告", 2)
    p.paragraph("8根A对应体积分数8V_A/L^3=0.0001130973，即0.01131%。若按百分号后保留两位，应报告0.01%；但必须同时保留“8根”和未过度舍入的0.01131%，因为相邻整数在两位百分数下可能显示相同。")

    p.heading("7  问题四：混合填充的整数成本优化")
    p.heading("7.1  优化模型", 2)
    p.paragraph("以n_A,n_B为整数决策变量，目标是最小化式(2)的材料成本，约束为总导通概率不低于0.90。由于总导通概率难以精确闭式计算，本文采用“候选方案用下界证明可行、所有更便宜方案用上界证明不可行”的双向证书。")
    p.equation("min C(n_A,n_B),   s.t. P(T)≥0.90,   n_A,n_B∈Z_≥0")
    p.paragraph("枚举域并非任意截取。只要先找到一个由直接贯通下界证明可行的候选(n_A*,n_B*)，其成本C*就是全局最优值的上界；任何更优解都必须满足c_A n_A+c_B n_B<C*，从而0≤n_A< C*/c_A、0≤n_B<C*/c_B，候选集合自动变为有限整数三角域。对其中每一点，若总导通上界仍小于0.90，则该点在未知真实概率下也必不可行。")
    p.code_block("先由直接贯通下界寻找一个可行候选，记成本 C*\nfor n_A = 0,...,floor(C*/c_A):\n    for n_B = 0,...,floor((C*−c_A n_A)/c_B):\n        若成本 < C*：计算 P_upper(n_A,n_B)\n        若 P_upper ≥ 0.90：保留为未排除候选\n若未排除集合为空，则当前候选为全局最优")
    p.heading("7.2  非负整数域的边界解", 2)
    p.paragraph("若允许某一类介质数量为零，0A+57B的直接贯通下界为0.9023976480，成本0.0955044元，故该方案可行。以该成本为上限，枚举全部216个更低成本非负整数点；其中总导通上界最大的是0A+56B，上界0.8984306753<0.90。因此所有更便宜方案均不可行，0A+57B在非负整数域严格最优。")
    p.heading("7.3  两类介质均为正的主口径", 2)
    p.paragraph("题目中的“同时填充A、B”可能要求两类介质均出现。增加n_A>=1,n_B>=1后，1A+50B的直接贯通下界为0.9031977272，成本0.0986198元。枚举164个更低成本正混合点后，最危险方案1A+49B的总导通上界为0.8992436792<0.90。因此若按“同时填充”解释，应把1A+50B作为正式答案，0A+57B仅作为放宽约束后的对照。")
    q4 = RESULTS["Q4"]
    q4_rows = [[f"{x['a_count']}A+{x['b_count']}B", f"{x['cost_cny']:.6f}", f"{x['direct_bridge_probability']:.6f}", f"{x['conduction_upper_bound']:.6f}", "更便宜，不可行"] for x in q4["cheaper_frontier"]]
    q4_rows += [["0A+57B", f"{q4['selected']['cost_cny']:.6f}", f"{q4['selected']['direct_bridge_lower_bound']:.6f}", "-", "非负整数域最优"], ["1A+50B", f"{q4['strictly_positive_mixture']['selected']['cost_cny']:.6f}", f"{q4['strictly_positive_mixture']['selected']['direct_bridge_lower_bound']:.6f}", "-", "正混合域最优"]]
    p.table("Q4低成本前沿与两种口径的候选解", ["方案", "成本/元", "直接下界", "总上界", "结论"], q4_rows, [2.5, 2.5, 3.0, 3.0, 5.2], font_size=8.8)
    p.figure("C08_q4_cost_frontier.png", "Q4成本—概率界前沿与候选解证书", 15.3, research=True)
    p.heading("7.4  成本效率解释", 2)
    p.paragraph("A单体直接贯通概率高，但单体成本约为B的8.86倍；B虽需要更多数量，却具有更高的单位成本概率收益，因此非负整数域最优点落在纯B边界。强制两类均出现时，加入1根A可以减少7个B，形成1A+50B的正混合最优点。该解释只针对本题给定价格和尺寸，若B单价上升或半径改变，整数前沿会发生跳变。")
    p.paragraph("以对数可靠性增益−log(1−q)除以单体成本作局部效率比较，B在当前价格下明显占优；但整数约束使替代关系不是连续比例，最优方案必须在离散格点上验证。1A+49B是主口径下最危险的更低成本点，其总上界仅比0.90低约7.56×10^-4，因此图形同时标出该点而不是只展示最优解。")

    p.heading("8  模型检验、灵敏度与稳健性")
    p.heading("8.1  几何算法检验", 2)
    p.paragraph("几何测试覆盖平行、相交、端点最近、轴段换序和多对批量计算等情形；对附件三组数据，全对全枚举与KD树广相位得到相同的连通结论和边数。正例路径逐边检查s,t、端面余量、轴距与正交残差，电极边另行保存表面间隙；负例则利用胶囊超集反证。当前测试共18项通过，其中概率边界测试明确区分无条件电极接触层a_i+g与非越界剩余薄壳g，防止把g/L误写成无条件接触概率。")
    p.table("计算与证据的验证矩阵", ["对象", "主计算", "独立复核", "通过标准", "结果"], [
        ["轴段最短距离", "候选分段闭式核", "换序、平行、相交与端点样例", "距离对称且退化样例符合解析值", "通过"],
        ["Q1候选边", "KD广相位+精算", "全对全O(n²)枚举", "三组边数与连通结论一致", "通过"],
        ["正例路径", "BFS父指针回溯", "逐边几何证书", "电极间隙≤g；内部边轴距≤61.8 nm", "通过"],
        ["Q3阈值", "直接下界与薄壳上界", "7、8根高精度重算", "7根上界<0.90<8根下界", "通过"],
        ["Q4最优性", "成本以下整数枚举", "记录最危险低价点", "所有更低成本点上界<0.90", "通过"],
    ], [2.8, 3.6, 4.1, 5.0, 1.8], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER], font_size=8.4)
    p.heading("8.2  概率与整数域复核", 2)
    p.paragraph("q_A由方向积分解析得到，q_B由球的边界层宽度直接得到；Q3分别重算7根上界与8根下界。Q4不是只比较表7中的前沿点，而是完整枚举成本严格低于候选的所有整数点，再取其总导通上界最大者。因此216和164是搜索域内实际候选数，不是抽样规模。")
    p.heading("8.3  几何参数灵敏度", 2)
    p.paragraph("为考察结构敏感性，在保持其余条件不变的设计情景中令A高度H取3500-6500 nm、B半径r_B取120-280 nm。该范围不是实测误差或制造分布，只用于观察解析阈值的整数跳变。H增大时q_A线性增大，达到90%所需A数量呈阶梯下降；B半径增大时q_B=2r_B/L增大，所需B数量下降。题设基准H=5000 nm对应8根A，r_B=200 nm对应57个B。")
    p.figure("C09_sensitivity_threshold_counts.png", "设计情景下几何尺寸对90%充分数量的影响", 15.3, research=True)
    p.heading("8.4  假设敏感性与失效情形", 2)
    p.table("关键假设变化对结论的影响", ["变化", "直接影响", "最可能受影响的结论", "建议改进"], [
        ["A取向偏向x轴", "q_A增大", "Q3阈值下降，Q4更偏向A", "用实测取向分布替代各向同性积分"],
        ["粒子不可重叠", "位置不再独立", "Q2-Q4乘积式和联合界需重建", "随机序列吸附或排斥点过程模拟"],
        ["粒子团聚", "局部连接增强但边界分布改变", "总概率与解析界间隙增大", "相关随机几何图与蒙特卡洛校准"],
        ["周期片段不保持导体身份", "直接贯通事件失效", "Q2-Q4全部数值失效", "按新边界物理重新定义连接图"],
        ["要求A、B均出现", "可行域删去坐标轴", "Q4主答案变为1A+50B", "论文同时报告两种口径"],
    ], [3.0, 4.0, 4.6, 4.6], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.8)
    p.heading("8.5  取向分布极端情景", 2)
    p.paragraph("各向同性不是题面给定事实，而是把“任意方向”转化为概率的补充模型。若所有A都与x轴平行，则单根直接贯通概率q_A=H/L=0.500；若所有轴都限制在yz平面，则仅端面圆盘在x方向提供支撑，q_A=2r_A/L=0.006。二者相差两个数量级，说明取向分布会直接改变Q2的风险数量级以及Q3、Q4的整数阈值。故0.25471239应理解为各向同性方向平均，而非所有取向都成立的材料常数。")

    p.heading("9  模型评价与改进")
    p.heading("9.1  模型优点", 2)
    p.paragraph("第一，Q1采用路径证书而非只给程序布尔值，负例和正例分别由超集反证与侧面充分证书支撑。第二，Q2-Q4明确区分精确直接贯通概率、总导通下界和总导通上界，避免把蒙特卡洛频率或下界误当作真实概率。第三，Q4对候选成本以下的有限整数域完整枚举，最优性证据可逐点复算。第四，所有关键数值集中写入结构化结果文件并由测试检查，论文图表直接从附件和结果文件生成。")
    p.heading("9.2  模型局限", 2)
    p.paragraph("随机模型的主要局限是题面未唯一指定概率测度，独立均匀、各向同性以及周期片段保持母体电学身份均需作为条件列明；允许重叠则来自题面。非直接通路联合上界只适合做不足性证明，不能精确刻画粒子间簇连通。Q1的路径核验只服务于所给附件结论，不代表候选图中的所有端部边都已完成通用平端圆柱实体距离分类。")
    p.heading("9.3  后续改进", 2)
    p.paragraph("后续可在三方面扩展：其一，实现平端圆柱—圆柱、圆柱—球的完整最近距离求解器，替代胶囊候选超图；其二，引入不可重叠随机序列吸附、取向偏置或团聚点过程，并用方差缩减蒙特卡洛估计真实导通概率；其三，在解析界与模拟之间建立校准区间，使模型既能证明阈值，又能给出阈值以外更精细的工程概率估计。")

    p.heading("10  结论")
    p.paragraph("(1) 附件组1不导通，组2和组3导通；组2显式路径为左面-2-12-24-39-右面，组3为左面-63-264-216-351-右面。所有正例介质间边均具有轴段内部最短点证书。", first=False)
    p.paragraph("(2) 在独立均匀中心、独立各向同性取向及周期片段保持母体电学身份的条件下，A体积分数0.50%、0.60%、0.70%和1.00%对应354、424、495和707根A；总不导通概率分别不超过10^-45.20、10^-54.13、10^-63.20和10^-90.27量级。", first=False)
    p.paragraph("(3) 在同一概率模型条件下，仅填充A时，7根的总导通上界0.872279低于90%，8根的直接贯通下界0.904810高于90%，故最小数量为8根；体积分数为0.01131%，按百分号后两位报告为0.01%。", first=False)
    p.paragraph("(4) 按“同时填充”要求A、B均出现时，主答案为1A+50B，成本0.09862元；若放宽为非负整数域，边界解为0A+57B，成本0.09550元。两者都受上述概率模型条件限制。", first=False)

    p.heading("参考文献")
    refs = [
        "[1] 华数杯大学生数学建模竞赛组委会，2026年第七届华数杯大学生数学建模竞赛A题：微构体中填充导电介质的仿真优化，2026。",
        "[2] 华数杯大学生数学建模竞赛组委会，2026年华数杯数学建模竞赛论文格式规范与提交说明，https://m.saikr.com/chinamcm26，访问日期：2026-08-13。",
        "[3] Feller W. An Introduction to Probability Theory and Its Applications, Vol. 1. New York: Wiley, 1968.",
        "[4] Ericson C. Real-Time Collision Detection. Boca Raton: CRC Press, 2005.",
        "[5] Meester R, Roy R. Continuum Percolation. Cambridge: Cambridge University Press, 1996.",
        "[6] Balberg I, Anderson C H, Alexander S, et al. Excluded volume and its relation to the onset of percolation. Physical Review B, 1984, 30(7): 3933. DOI: 10.1103/PhysRevB.30.3933.",
        "[7] Foygel M, Morris R D, Anez D, et al. Theoretical and computational studies of carbon nanotube composites and suspensions: Electrical and thermal conductivity. Physical Review B, 2005, 71: 104201. DOI: 10.1103/PhysRevB.71.104201.",
        "[8] Otten R H J, van der Schoot P. Connectivity percolation of polydisperse anisotropic nanofillers. Journal of Chemical Physics, 2011, 134: 094902. DOI: 10.1063/1.3559004.",
        "[9] Chatterjee A P, Grimaldi C. Random geometric graph description of connectedness percolation in rod systems[J]. Physical Review E, 2015, 92: 032121. DOI: 10.1103/PhysRevE.92.032121.",
        "[10] Xu W, Su X, Jiao Y. Continuum percolation of congruent overlapping spherocylinders[J]. Physical Review E, 2016, 94: 032122. DOI: 10.1103/PhysRevE.94.032122.",
        "[11] Schilling T, Miller M, van der Schoot P. Percolation in suspensions of hard nanoparticles: from spheres to needles[J]. Europhysics Letters, 2015, 111: 56004. DOI: 10.1209/0295-5075/111/56004.",
    ]
    for ref in refs:
        p.paragraph(ref, first=False, size=10.5, after=2)

    p.page_break()
    p.heading("附录A  复现环境与命令")
    p.paragraph("复现环境：Python 3.13；numpy 2.2.6；pandas 3.0.3；openpyxl 3.1.5；scipy 1.16.3；pytest 9.1.0。", first=False)
    commands = """cd mathmodel/runs/huashubei-2026-final-001
python -m pip install -r requirements.txt
python src/a/build_corrected_results.py
python src/a/build_research_figures.py
python src/a/build_full_paper.py
python -m pytest tests -q"""
    p.code_block(commands)
    p.paragraph("当前A题几何与解析模型测试共18项通过；正文、图表和机器结果均可由上述命令重建。", first=False)

    p.heading("附录B  证据文件索引")
    p.paragraph("机器可读结果位于outputs/data/final_results.json；核心算法位于src/a/analytic_bounds.py、src/a/geometry.py和src/a/build_corrected_results.py；17组图形候选位于outputs/figure_candidates，12幅入文图的PNG、PDF与SVG源位于outputs/figures_research。为控制论文篇幅，完整源代码不再嵌入正文，而随仓库一并提交。", first=False)

    docx = OUT / f"{OUTPUT_STEM}.docx"
    md = OUT / f"{OUTPUT_STEM}.md"
    p.save(docx, md)
    return docx, md


def build_paper_condensed_archive():
    """Retained compact variant for comparison; the complete paper is built by main()."""
    OUT.mkdir(parents=True, exist_ok=True)
    p = Paper()
    p.meta_line()
    p.title()
    p.abstract_heading()
    p.paragraph(
        "本文研究边长10000 nm的立方微构体中导电介质的左右连通判定、随机填充概率和材料成本优化。"
        "确定性问题以附件每行给出的平端圆柱轴段为一个介质节点，用胶囊超图生成保守候选边，再以轴段最近点的内部性与正交性核验导通路径；"
        "随机问题则先明确中心独立均匀、圆柱轴向独立各向同性以及周期平移片段保持同一介质身份三项生成机制，"
        "以单粒子跨界直接贯通事件给出总导通概率下界，并以两侧终端薄壳的有序对联合界给出总导通概率上界。"
    )
    p.paragraph(
        "按上述口径，附件组1不导通，组2和组3导通，显式路径分别为左面-2-12-24-39-右面与左面-63-264-216-351-右面。"
        "当A的目标体积分数为0.50%、0.60%、0.70%和1.00%时，取整后分别为354、424、495和707根，"
        "总不导通概率至多为10⁻⁴⁵·²⁰、10⁻⁵⁴·¹³、10⁻⁶³·²⁰和10⁻⁹⁰·²⁷量级；因此在常用报告精度内均为1，但数学上仍严格小于1。"
    )
    p.paragraph(
        "仅填充A时，7根的总导通概率上界为0.872279，而8根的直接贯通下界为0.904810，故最低数量为8根，"
        "实际体积分数为0.01131%，按题意保留百分号后两位为0.01%。按“同时填充”要求两类介质均出现，"
        "最低成本方案为1A+50B，成本0.09862元；若允许某类数量为零，放宽域对照为0A+57B，成本0.09550元。"
        "全部阈值和最优性结论均由可行侧下界与不可行侧上界共同闭合。"
    )
    p.keyword("周期截断；随机几何图；概率夹逼；联合界；整数优化")
    p.page_break()

    p.heading("1  问题分析与研究路线")
    p.heading("1.1  四问的任务合同", 2)
    p.paragraph(
        "四个问题并非彼此独立：问题一建立确定性接触图，问题二明确随机生成机制，问题三在同一概率模型上寻找整数阈值，"
        "问题四再叠加材料成本和正整数约束。为避免只给结论而不说明证据类型，表1把每一问的输入、输出和验收条件并列给出。"
    )
    p.table("四个问题的输入、输出与核心证据", ["问题", "输入", "输出", "核心证据"], [
        ["问题一", "三组A轴段坐标", "每组是否导通及显式路径", "胶囊候选图、BFS路径、逐边几何证书"],
        ["问题二", "A体积分数与随机生成机制", "四个体积分数下的概率界", "粒子数换算、直接贯通下界、失败概率上界"],
        ["问题三", "仅A、目标概率0.90", "最低填充量", "8根下界充分、7根上界不足、单调性"],
        ["问题四", "A/B成本与目标概率0.90", "最低成本整数方案", "候选下界可行、全部更便宜点上界不可行"],
    ], [1.8, 4.4, 4.4, 6.1], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=9)

    p.heading("1.2  可识别口径与基线选择", 2)
    p.paragraph(
        "附件说明“每一行表示一个介质A”，但每行只给两个轴端点，无法额外表达一个被多个边界切开的全部片段。"
        "因此问题一把附件行本身视为可观测的盒内轴段，不跨行猜测同一介质，也不依据5000 nm理论长度补造未给坐标。"
        "这是对附件可识别信息的最小使用口径；若官方另给片段配对或原始中心—方向数据，接触图应据此重建。"
    )
    p.paragraph(
        "随机问题的朴素基线是逐次生成完整几何构型并做蒙特卡洛统计。但问题三要求证明“7不足、8充分”，问题四要求排除所有更便宜整数点，"
        "有限样本频率难以提供严格证书。本文因此采用解析概率夹逼作为主模型：下界负责证明可行，上界负责证明不足；"
        "蒙特卡洛只适合作为未来的区间收紧手段，而不冒充当前结论来源。"
    )
    p.heading("1.3  总体技术路线", 2)
    p.paragraph(
        "全文遵循“数据审计—几何图判定—概率空间声明—上下界推导—整数域穷举—独立复核”的顺序。"
        "每个强结论同时说明适用条件、计算值和反例边界；图表只承担空间路径、数量级、阈值夹逼和成本前沿四类证据功能。"
    )

    p.heading("2  模型条件、符号与量纲")
    p.heading("2.1  题面条件与补充随机机制", 2)
    p.table("模型条件、使用位置与失效后果", ["性质", "条件或解释", "使用位置", "失效后果"], [
        ["题面条件", "随机介质允许贯穿、重叠", "问题二至四", "增加排斥后独立乘积式需重建"],
        ["附件口径", "问题一以每行给定轴段为一个A", "问题一", "若存在额外片段配对，接触图需更新"],
        ["随机机制", "中心相互独立且在立方体内均匀", "问题二至四", "团聚或排斥会改变概率界"],
        ["随机机制", "A轴向相互独立且各向同性", "问题二至四", "取向偏置会改变单粒子跨界概率"],
        ["周期解释", "平移片段与母体保持同一介质身份", "问题二至四", "若电学断开，直接贯通事件失效"],
    ], [2.2, 6.4, 2.4, 5.7], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9)
    p.paragraph(
        "题面“任意位置、任意方向”只给出允许域，没有唯一指定概率分布。独立均匀与各向同性是最大对称性的无信息基线，"
        "不是材料固有属性。后文所有概率、最小数量和最低成本均条件于表2所列随机机制。"
    )

    p.heading("2.2  主要符号", 2)
    p.table("主要符号及定义", ["符号", "定义", "单位或定义域"], [
        ["L, g", "立方体边长、几何导通阈值", "10000 nm，1.8 nm"],
        ["H, r_A", "A的高度、半径", "5000 nm，30 nm"],
        ["r_B", "B的半径", "200 nm"],
        ["n_A, n_B", "A、B的整数数量", "非负整数；主口径均≥1"],
        ["a_x", "粒子在x方向的支撑半宽", "nm"],
        ["D, T", "至少一个粒子直接贯通、整体左右导通", "事件"],
        ["Sᵢᴸ, Sᵢᴿ", "粒子i未越界但接触左/右电极的薄壳事件", "事件"],
        ["P_dir, C", "直接贯通概率、材料成本", "[0,1]，元"],
    ], [2.5, 8.2, 5.7], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER], font_size=9.2)

    p.heading("2.3  体积与成本换算", 2)
    p.equation("V_A=πr_A²H=1.4137167×10⁷ nm³，V_B=4πr_B³/3=3.3510322×10⁷ nm³",
               md_expression=r"V_A=\pi r_A^2H=1.4137167\times10^7\,\mathrm{nm}^3,\quad V_B=\frac{4\pi r_B^3}{3}=3.3510322\times10^7\,\mathrm{nm}^3")
    p.equation("C=1.05V_An_A/10⁹+0.05V_Bn_B/10⁹=0.0148440253n_A+0.0016755161n_B",
               md_expression=r"C=\frac{1.05V_An_A+0.05V_Bn_B}{10^9}=0.0148440253n_A+0.0016755161n_B")
    p.paragraph(
        "式(1)按平端圆柱和球的真实体积计算，式(2)把nm³换算为μm³。所有优化先在整数域内完成，"
        "体积分数只用于结果报告，不用过早四舍五入的百分数反推数量。"
    )

    p.heading("3  问题一：确定性几何连通判定")
    p.heading("3.1  附件审计与行级接触图", 2)
    p.paragraph(
        "三组附件分别包含12、49和535行A数据，轴段长度范围依次为1019.117-5000.000 nm、577.474-5000.000 nm和525.697-5000.000 nm。"
        "不少端点恰位于盒边界，说明表中保留了截断后的轴段表达。本文建立无向图G=(V,E)：每行对应一个介质节点，另设左右电极节点；"
        "介质—介质或介质—电极的表面最短距离不超过g时连边，左右电极在同一连通分量中当且仅当该组导通。"
    )
    p.equation("(i,j)∈E ⇔ dist(Aᵢ,Aⱼ)≤g；(L,i)∈E ⇔ dist(Aᵢ,F_L)≤g",
               md_expression=r"(i,j)\in E\iff \operatorname{dist}(A_i,A_j)\le g,\qquad (\mathrm{L},i)\in E\iff \operatorname{dist}(A_i,F_L)\le g")

    p.heading("3.2  胶囊候选边与平端圆柱证书", 2)
    p.paragraph(
        "把轴段向外膨胀半径r_A得到胶囊体，胶囊包含相同轴段的平端圆柱。故两轴段距离d_axis≤2r_A+g=61.8 nm可生成不漏真的候选边；"
        "但端部半球会带来假阳性，所以候选边数不能直接称为真实平端圆柱边数。"
    )
    p.paragraph(
        "对导通路径上的每条候选边，设两轴段的最近点参数为s,t。若0<s<1、0<t<1，凸二次距离的一阶最优条件保证最近点连线同时垂直于两条轴线。"
        "沿该连线分别在两个截面圆盘内移动半径r_A，不改变轴向参数，也不会越过端面；因此实体表面间隙为max(0,d_axis-2r_A)。"
        "这给出了平端圆柱侧面接触的充分证书。数值实现同时保存端面最小余量和正交残差，避免仅凭图片判定。"
    )
    p.figure("C03_q1_path_certificate.png",
             "组2与组3显式通路的三维几何证书。红线为路径介质，绿虚线为相邻轴段的内部最近点连线，标注值为轴距/nm；线宽仅用于辨识，不代表实际半径。",
             15.8, research=True)

    p.heading("3.3  三组构型结果", 2)
    p.table("三组附件构型的连通结果", ["组别", "节点数", "候选边数", "左接触", "右接触", "结论", "显式路径"],
            q1_result_rows(), [1.1, 1.4, 1.7, 1.3, 1.3, 1.4, 7.7],
            aligns=[WD_ALIGN_PARAGRAPH.CENTER]*6+[WD_ALIGN_PARAGRAPH.LEFT], font_size=8.6)
    p.paragraph(
        "组1在更易连通的胶囊候选超图中仍不连通，因此按附件行级口径，真实平端圆柱图必不连通。组2、组3分别恢复出4个介质节点组成的通路。"
        "表5中6条介质间边的最近点均在轴段内部，端面最小余量为正，正交残差接近机器精度，因而不是胶囊端部造成的假接触。"
    )
    p.table("正例路径的平端圆柱侧面接触证书", ["组别", "边", "轴距/nm", "s", "t", "端面余量/nm", "正交残差", "结论"],
            q1_cert_rows(), [1.1, 1.7, 2.0, 1.1, 1.1, 2.5, 2.4, 1.4], font_size=8.4)

    p.heading("4  问题二至四的统一随机概率模型")
    p.heading("4.1  单粒子跨界直接贯通", 2)
    p.paragraph(
        "题面把越界部分平移到相对侧并仍称为同一介质。本文据此把平移片段视为同一导体的周期表示："
        "只要某粒子沿x方向越过任一边界，它在左右两侧的两个片段属于同一介质，故形成左右直接贯通。"
        "记至少一个粒子发生该事件为D，整体导通为T，则D⊆T。若采用“平移片段电学断开”的另一物理解释，本章及问题二至四的数值需全部重建。"
    )
    p.heading("4.2  A与B的单粒子概率", 2)
    p.paragraph(
        "设A轴向单位向量的x分量为u_x。平端圆柱在x方向的支撑半宽由轴向投影和端面圆盘投影共同组成："
    )
    p.equation("a_x=(H/2)|u_x|+r_A√(1-u_x²)",
               md_expression=r"a_x=\frac{H}{2}\lvert u_x\rvert+r_A\sqrt{1-u_x^2}")
    p.paragraph(
        "固定方向时，中心位于两侧合计宽度2a_x的边界层即越界。各向同性无向轴满足E|u_x|=1/2和E√(1-u_x²)=π/4，故："
    )
    p.equation("q_A=2E(a_x)/L=(H/2+πr_A/2)/L=0.25471238898",
               md_expression=r"q_A=\frac{2\mathbb E(a_x)}{L}=\frac{H/2+\pi r_A/2}{L}=0.25471238898")
    p.equation("q_B=2r_B/L=0.04", md_expression=r"q_B=\frac{2r_B}{L}=0.04")
    p.paragraph("独立粒子中至少一个直接贯通的概率由补事件乘积得到：")
    p.equation("P_dir=P(D)=1-(1-q_A)ⁿᴬ(1-q_B)ⁿᴮ",
               md_expression=r"P_{\mathrm{dir}}=P(D)=1-(1-q_A)^{n_A}(1-q_B)^{n_B}")

    p.heading("4.3  非直接通路的终端薄壳上界", 2)
    p.paragraph(
        "为证明较低数量或较低成本方案不足，必须给总导通概率上界。固定粒子方向后，若粒子不越界却仍接触左电极，"
        "其中心只能落在宽度恰为g的剩余薄壳；右侧同理。该宽度与粒子支撑半宽无关。"
    )
    p.equation("Sᵢᴸ=Cᵢᴸ∩Dᵢᶜ，P(Sᵢᴸ)=g/L",
               md_expression=r"S_i^L=C_i^L\cap D_i^c,\qquad P(S_i^L)=\frac{g}{L}")
    p.paragraph(
        "在D未发生而T发生时，通路的左右终端必须由不同粒子i≠j承担；否则同一粒子同时接触两端，已经属于D。"
        "因此有集合包含T∩Dᶜ⊆⋃_{i≠j}(Sᵢᴸ∩Sⱼᴿ)。跨粒子中心独立，对n(n-1)个有序对使用Boole联合界[2]，得到："
    )
    p.equation("P_dir≤P(T)≤min{1，P_dir+n(n-1)(g/L)²}，n=n_A+n_B",
               md_expression=r"P_{\mathrm{dir}}\le P(T)\le \min\left\{1,\ P_{\mathrm{dir}}+n(n-1)\left(\frac{g}{L}\right)^2\right\},\quad n=n_A+n_B")
    p.paragraph(
        "上界忽略中间粒子是否真的能连接，因而可能偏松，但不会漏掉任何曲折通路；它只用于证明“不足”，不作为真实概率点估计。"
        "下界只由充分事件构成，也不能写成总概率真值。"
    )

    p.heading("5  问题二：给定体积分数下的导通概率")
    p.heading("5.1  体积分数到整数数量", 2)
    p.paragraph("按题意将体积分数对应的A数量四舍五入为最近整数，并同时报告取整后的实际体积分数：")
    p.equation("n_A=round(φL³/V_A)，φ_actual=n_AV_A/L³",
               md_expression=r"n_A=\operatorname{round}\left(\frac{\phi L^3}{V_A}\right),\qquad \phi_{\mathrm{actual}}=\frac{n_AV_A}{L^3}")
    q2_rows = [[f"{100*r['requested_fraction']:.2f}%", r["a_count"], f"{100*r['achieved_fraction']:.5f}%",
                f"{r['log10_failure_probability_upper_bound']:.2f}",
                f"≥1-10^({r['log10_failure_probability_upper_bound']:.2f})"] for r in RESULTS["Q2"]]
    p.table("四个体积分数下的总导通下界与总不导通上界", ["目标体积分数", "A数量", "实际体积分数", "log₁₀不导通上界", "总导通概率下界"],
            q2_rows, [2.4, 1.6, 2.7, 4.0, 5.2], font_size=9)
    p.figure("C05_q2_failure_lollipop.png",
             "四个离散工况的总不导通概率上界。纵轴为log₁₀数量级，点旁同时标出取整后的A数量和概率上界；这些量来自解析补事件，不是蒙特卡洛频率。",
             15.3, research=True)
    p.paragraph(
        "粒子间进一步接触只会提高导通率，所以图2与表6给出的失败概率确为总不导通概率的上界。"
        "二进制浮点无法表示1-10⁻⁴⁵这一差异，机器结果保存对数失败上界和符号表达式，而不把下界误存为1.0。"
    )

    p.heading("6  问题三：仅填充A的最低填充量")
    p.heading("6.1  候选阈值与相邻整数证书", 2)
    p.paragraph("由直接贯通下界定位候选数量：")
    p.equation("n_A≥⌈log(0.10)/log(1-q_A)⌉=8",
               md_expression=r"n_A\ge\left\lceil\frac{\log(0.10)}{\log(1-q_A)}\right\rceil=8")
    q3_rows = []
    for r in RESULTS["Q3"]["proof_rows"][-2:]:
        q3_rows.append([r["a_count"], f"{r['direct_bridge_lower_bound']:.9f}",
                        f"{r['non_direct_path_upper_addition']:.2e}", f"{r['conduction_upper_bound']:.9f}",
                        "不足" if r["a_count"] == 7 else "充分"])
    p.table("7根不足与8根充分的概率夹逼", ["A数量", "直接贯通下界", "非直接上界增量", "总导通上界", "阈值结论"],
            q3_rows, [1.6, 3.8, 4.0, 3.8, 2.5], font_size=9.2)
    p.figure("C04_q3_bounds_band.png",
             "A数量1-8时的解析概率界及7、8根局部放大。蓝线是直接贯通下界，红虚线是总导通上界，橙色窄带是两者的解析差而非置信区间。",
             15.3, research=True)
    p.paragraph(
        "8根时下界0.904810024>0.90，故8根充分；7根时上界0.872278889<0.90，故7根不足。"
        "在同一无限随机序列上把前n根作为n粒子构型，加入第n+1根不会删除已有路径，所以真实导通概率关于n非降；"
        "因此7根上界同时排除了0-7根，8根即为严格最小整数。"
    )
    p.heading("6.2  填充率报告", 2)
    p.paragraph(
        "8根A的实际体积分数为8V_A/L³=0.0001130973，即0.01131%。按题意精确到百分号后两位应写0.01%，"
        "但正文同时保留8根和0.01131%，避免相邻整数在两位百分数下显示相同。"
    )

    p.heading("7  问题四：混合填充的整数成本优化")
    p.heading("7.1  主口径与证明方法", 2)
    p.paragraph(
        "题目使用“同时填充介质A和B”，本文把两类数量均为正作为主口径；允许某类为零的非负整数域只作对照。"
        "主问题为："
    )
    p.equation("min C(n_A,n_B)，s.t. P(T)≥0.90，n_A,n_B∈Z，n_A≥1，n_B≥1",
               md_expression=r"\min C(n_A,n_B)\quad\mathrm{s.t.}\quad P(T)\ge0.90,\quad n_A,n_B\in\mathbb Z,\ n_A\ge1,\ n_B\ge1")
    p.paragraph(
        "对任一候选，用式(7)的直接贯通下界证明可行；再枚举其成本以下的全部整数点，用式(9)的总上界逐点排除。"
        "这不是启发式搜索，候选下方的有限整数域全部被检查。"
    )
    q4 = RESULTS["Q4"]
    bad_rel = q4["maximum_upper_bound_among_cheaper"]
    rel = q4["selected"]
    bad_pos = q4["strictly_positive_mixture"]["maximum_upper_bound_among_cheaper"]
    pos = q4["strictly_positive_mixture"]["selected"]
    q4_rows = [
        [f"{bad_pos['a_count']}A+{bad_pos['b_count']}B", f"{bad_pos['cost_cny']:.6f}", "-", f"{bad_pos['conduction_upper_bound']:.9f}", "主口径最危险的更便宜点，不可行"],
        [f"{pos['a_count']}A+{pos['b_count']}B", f"{pos['cost_cny']:.6f}", f"{pos['direct_bridge_lower_bound']:.9f}", "-", "主口径最低成本，可行"],
        [f"{bad_rel['a_count']}A+{bad_rel['b_count']}B", f"{bad_rel['cost_cny']:.6f}", "-", f"{bad_rel['conduction_upper_bound']:.9f}", "放宽域最危险的更便宜点，不可行"],
        [f"{rel['a_count']}A+{rel['b_count']}B", f"{rel['cost_cny']:.6f}", f"{rel['direct_bridge_lower_bound']:.9f}", "-", "放宽域最低成本，可行"],
    ]
    p.table("主口径与放宽域的可行—不可行双向证书", ["方案", "成本/元", "直接下界", "总上界", "证明角色"],
            q4_rows, [2.1, 2.3, 3.1, 3.1, 6.1],
            aligns=[WD_ALIGN_PARAGRAPH.CENTER]*4+[WD_ALIGN_PARAGRAPH.LEFT], font_size=8.8)
    p.figure("C08_q4_cost_frontier.png",
             "成本阈值以下的整数候选与两种口径证书。空心圆为更低成本正混合点的总上界，橙色菱形为主口径1A+50B的直接下界；绿色星号为放宽域0A+57B的直接下界。",
             15.3, research=True)
    p.paragraph(
        "主口径中，1A+50B的下界为0.903197727，成本0.0986198元；164个更便宜正混合点中，上界最大的是1A+49B，"
        "其值0.899243679仍低于0.90，故1A+50B为条件全局最优。放宽域中，0A+57B的下界为0.902397648，成本0.0955044元；"
        "216个更便宜点全部被排除，其中0A+56B的上界最大，为0.898430675。"
    )
    p.heading("7.2  成本效率解释", 2)
    p.paragraph(
        "A的单粒子直接贯通概率较高，但单体成本约为B的8.86倍。按本题尺寸和单价，B具有更高的单位成本跨界收益，"
        "故放宽域最优点落在纯B边界；主口径强制两类都出现后，加入1根A可减少7个B，形成1A+50B。"
        "这一排序只对题设尺寸、价格和随机机制成立，不应外推为材料性能的一般结论。"
    )

    p.heading("8  模型检验与稳健性")
    p.heading("8.1  计算与证据链复核", 2)
    p.paragraph(
        "几何实现同时采用全对全枚举和KD树广相位，三组的候选边数、接触数和连通结论一致；路径边进一步检查最近点参数、端面余量和正交残差。"
        "概率部分分别由解析积分、补事件乘积和联合界独立计算；整数优化不只展示前沿，而是重新计算候选成本以下每个整数点的约束上界。"
        "所有表格与数据图均从final_results.json和附件生成，避免手工抄数造成正文—图表漂移。"
    )
    p.heading("8.2  取向极限与结构敏感性", 2)
    p.paragraph(
        "各向同性是最关键的补充假设。若A全部平行于x轴，则q_A=H/L=0.500，使直接贯通下界达到90%只需4根；"
        "各向同性时需8根；若A轴全部位于yz平面，仅端面半径贡献x向支撑，q_A=2r_A/L=0.006，充分数量增至383根。"
        "三个极限相差近两个数量级，说明8根不是脱离取向分布的材料常数。"
    )
    p.table("取向情景对A直接贯通充分数量的影响", ["取向情景", "q_A", "使直接贯通下界≥0.90的数量", "解释"], [
        ["全部平行x轴", "0.500000", "4", "轴向投影最大"],
        ["独立各向同性", "0.254712", "8", "本文主模型"],
        ["全部位于yz平面", "0.006000", "383", "仅端面半径贡献x向支撑"],
    ], [3.2, 2.4, 5.0, 5.0], aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT], font_size=9)
    p.heading("8.3  失效情形", 2)
    p.paragraph(
        "若粒子不可重叠、发生团聚或取向相关，跨粒子独立性失效，式(7)与式(9)需用相应点过程重建；若周期平移片段不保持电连接，"
        "单粒子跨界不再是充分事件，问题二至四的数值全部失效；若目标是有效电导率而不是题设二值导通，还需引入材料电阻、接触电阻和隧穿电阻网络。"
    )

    p.heading("9  模型评价与改进")
    p.heading("9.1  模型优点", 2)
    p.paragraph(
        "模型的核心优势是证据与结论配对：问题一给显式路径和逐边几何证书；问题二保留极小失败概率的对数表达；"
        "问题三同时给出7根不足与8根充分；问题四对候选成本以下的有限整数域完整排除。"
        "解析界使阈值和最优性不依赖随机种子，也明确区分事件概率、总概率下界和总概率上界。"
    )
    p.heading("9.2  模型局限", 2)
    p.paragraph(
        "题面未唯一给出随机概率测度，因而随机问题只能是条件结论。终端薄壳上界用于证明不足但通常不紧，不能精确刻画中间粒子簇。"
        "问题一只对附件给定行级轴段作判定，没有在缺少片段配对信息时补造周期互补段；候选图中的非路径边也未逐条完成通用平端实体距离分类。"
    )
    p.heading("9.3  可执行改进", 2)
    p.paragraph(
        "后续可依次完成三项增强：首先取得原始中心—方向或片段ID，重建完整周期构型；其次实现平端圆柱、球和周期像的统一凸距离核；"
        "最后在解析界之外运行带方差缩减的有限盒蒙特卡洛，报告置信区间并校准上下界间隙。"
        "这些增强用于缩小不确定性，不改变当前论文必须显著声明条件的原则。"
    )

    p.heading("10  结论")
    p.paragraph("(1) 按附件每行给定轴段为一个A的可识别口径，组1不导通，组2和组3导通；两条显式路径分别为左面-2-12-24-39-右面和左面-63-264-216-351-右面，路径边通过内部最近点、端面余量和正交性核验。", first=False)
    p.paragraph("(2) 条件于中心独立均匀、A轴向独立各向同性及周期平移片段保持同一介质身份，目标体积分数0.50%、0.60%、0.70%和1.00%分别对应354、424、495和707根A，总不导通概率上界依次为10⁻⁴⁵·²⁰、10⁻⁵⁴·¹³、10⁻⁶³·²⁰和10⁻⁹⁰·²⁷量级。", first=False)
    p.paragraph("(3) 在同一随机机制下，7根A的总导通上界0.872279低于0.90，8根A的直接贯通下界0.904810高于0.90，故最低为8根；实际体积分数0.01131%，按题意报告为0.01%。", first=False)
    p.paragraph("(4) 按两类介质均出现的主口径，最低成本方案为1A+50B，成本0.09862元；若允许某类为零，放宽域对照为0A+57B，成本0.09550元。两项最优性均在声明的概率、周期身份、整数数量和线性实际体积成本条件下成立。", first=False)

    p.heading("参考文献")
    refs = [
        "[1] 华数杯大学生数学建模竞赛组委会. 2026年第七届华数杯大学生数学建模竞赛A题：微构体中填充导电介质的仿真优化[Z]. 2026.",
        "[2] Feller W. An Introduction to Probability Theory and Its Applications, Vol. 1[M]. New York: Wiley, 1968.",
        "[3] Ericson C. Real-Time Collision Detection[M]. Boca Raton: CRC Press, 2005.",
        "[4] Meester R, Roy R. Continuum Percolation[M]. Cambridge: Cambridge University Press, 1996.",
        "[5] Balberg I, Anderson C H, Alexander S, et al. Excluded volume and its relation to the onset of percolation[J]. Physical Review B, 1984, 30(7): 3933. DOI: 10.1103/PhysRevB.30.3933.",
        "[6] Otten R H J, van der Schoot P. Connectivity percolation of polydisperse anisotropic nanofillers[J]. Journal of Chemical Physics, 2011, 134: 094902. DOI: 10.1063/1.3559004.",
    ]
    for ref in refs:
        p.paragraph(ref, first=False, size=10.5, after=2)

    p.heading("附录A  复现入口")
    p.paragraph("关键结果由附件、解析公式与整数枚举脚本生成；运行环境和命令如下。", first=False)
    commands = """cd mathmodel/runs/huashubei-2026-final-001
python src/a/build_corrected_results.py
python src/a/build_research_figures.py
python src/a/build_full_paper.py
python -m pytest tests -q"""
    p.code_block(commands)
    p.heading("附录B  证据文件索引")
    p.paragraph(
        "机器可读结果位于outputs/data/final_results.json；几何与概率算法位于src/a/geometry.py和src/a/analytic_bounds.py；"
        "图形源文件保存在outputs/figures_research，均同时提供PNG、PDF和SVG；攻击性审查与修改理由保存在reports目录。"
        "完整源代码和测试随仓库提交，正文不重复粘贴。", first=False)

    docx = OUT / f"{OUTPUT_STEM}.docx"
    md = OUT / f"{OUTPUT_STEM}.md"
    p.save(docx, md)
    return docx, md


def main():
    docx, md = build_paper_legacy()
    print(docx)
    print(md)


if __name__ == "__main__":
    main()
